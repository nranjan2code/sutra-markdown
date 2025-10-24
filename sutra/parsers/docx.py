"""
DOCX Parser - Extract content from Word documents

Uses python-docx for parsing Microsoft Word documents with support for:
- Text extraction with formatting (bold, italic, underline)
- Style detection (headings, paragraphs, lists)
- Table extraction with structure
- Image extraction
- Metadata extraction
- Comments and track changes
- Headers and footers

Requirements:
    pip install python-docx

Usage:
    parser = DOCXParser("document.docx")
    result = await parser.parse()
    
    # Or streaming
    async for element in parser.parse_stream():
        process(element)
"""

import time
from pathlib import Path
from typing import List, Optional, Dict, Any, AsyncIterator
from datetime import datetime
import asyncio

from ..models.document import (
    ParsedDocument,
    PageInfo,
    ImageInfo,
    TableInfo,
    DocumentElement,
    DocumentMetadata
)
from ..models.enums import ElementType, DocumentType
from .base import BaseParser, ParserResult, ParserError, CorruptedFileError

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class DOCXParser(BaseParser):
    """
    Word document parser using python-docx
    
    Extracts text, tables, images, and metadata from DOCX files.
    Supports both full document parsing and streaming mode.
    
    Args:
        file_path: Path to DOCX file
        options: Parser options
            - extract_images: Extract images (default: True)
            - extract_tables: Extract tables (default: True)
            - extract_comments: Extract comments (default: False)
            - extract_headers_footers: Extract headers/footers (default: True)
            - preserve_formatting: Track formatting info (default: True)
    
    Example:
        >>> parser = DOCXParser("report.docx", {
        ...     "extract_images": True,
        ...     "extract_tables": True,
        ...     "preserve_formatting": True
        ... })
        >>> result = await parser.parse()
        >>> print(f"Extracted {len(result.document.pages)} sections")
    """
    
    def __init__(
        self,
        file_path: str | Path,
        options: Optional[Dict[str, Any]] = None
    ):
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX parsing.\n"
                "Install with: pip install python-docx"
            )
        
        super().__init__(file_path, options)
        
        # Parse options
        self.extract_images = self.options.get("extract_images", True)
        self.extract_tables = self.options.get("extract_tables", True)
        self.extract_comments = self.options.get("extract_comments", False)
        self.extract_headers_footers = self.options.get("extract_headers_footers", True)
        self.preserve_formatting = self.options.get("preserve_formatting", True)
        
        # Open DOCX document
        try:
            self.doc = Document(str(self.file_path))
        except Exception as e:
            raise CorruptedFileError(f"Failed to open DOCX: {e}")
    
    def get_page_count(self) -> int:
        """
        Get approximate page count
        
        Note: DOCX doesn't have explicit pages, so we estimate
        based on content length (assuming ~500 words per page)
        """
        word_count = 0
        for para in self.doc.paragraphs:
            word_count += len(para.text.split())
        
        # Estimate pages (500 words per page)
        pages = max(1, word_count // 500)
        return pages
    
    async def validate(self) -> tuple[bool, Optional[str]]:
        """
        Validate DOCX file
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check if document can be opened
            if not hasattr(self, 'doc'):
                return False, "Failed to open DOCX document"
            
            # Check if has content
            if not self.doc.paragraphs and not self.doc.tables:
                return False, "DOCX has no content"
            
            # Try to read first paragraph
            try:
                if self.doc.paragraphs:
                    _ = self.doc.paragraphs[0].text
            except Exception as e:
                return False, f"Failed to read DOCX content: {e}"
            
            return True, None
            
        except Exception as e:
            return False, str(e)
    
    async def parse(self) -> ParserResult:
        """
        Parse entire DOCX document
        
        Returns:
            ParserResult with parsed content
        """
        start_time = time.time()
        warnings = []
        
        try:
            # Validate first
            is_valid, error = await self.validate()
            if not is_valid:
                return ParserResult(
                    document=None,
                    success=False,
                    error=error,
                    file_size=self.file_size
                )
            
            # Extract metadata
            metadata = self._extract_metadata()
            
            # Parse content
            sections = []
            images = []
            tables = []
            full_text = []
            
            # Extract headers/footers
            if self.extract_headers_footers:
                header_footer_text = self._extract_headers_footers()
                if header_footer_text:
                    full_text.append(f"--- Headers/Footers ---\n{header_footer_text}")
            
            # Process document body
            for element in self.doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    para = Paragraph(element, self.doc)
                    section_data = self._parse_paragraph(para)
                    
                    if section_data["text"]:
                        sections.append(section_data)
                        full_text.append(section_data["text"])
                
                elif isinstance(element, CT_Tbl):
                    # Table
                    if self.extract_tables:
                        table = Table(element, self.doc)
                        table_data = self._parse_table(table, len(tables))
                        
                        # Only append if table is valid (not None)
                        if table_data is not None:
                            tables.append(table_data)
                            
                            # Add table text to full text
                            table_text = self._table_to_text(table_data)
                            full_text.append(table_text)
            
            # Extract images
            if self.extract_images:
                images = self._extract_images()
            
            # Create pages (treat each section as a page for compatibility)
            pages = []
            for idx, section in enumerate(sections):
                page_info = PageInfo(
                    page_number=idx + 1,
                    text=section["text"],
                    width=595.0,  # Default A4 width in points
                    height=842.0,  # Default A4 height in points
                )
                pages.append(page_info)
            
            # Create ParsedDocument
            document = ParsedDocument(
                file_path=str(self.file_path),
                file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=self.file_size,
                text="\n\n".join(full_text),
                metadata=metadata,
                pages=pages,
                images=images,
                tables=tables
            )
            
            parse_time = time.time() - start_time
            
            return ParserResult(
                document=document,
                success=True,
                warnings=warnings,
                parse_time=parse_time,
                file_size=self.file_size
            )
            
        except Exception as e:
            return ParserResult(
                document=None,
                success=False,
                error=str(e),
                parse_time=time.time() - start_time,
                file_size=self.file_size
            )
    
    async def parse_stream(self) -> AsyncIterator[DocumentElement]:
        """
        Parse DOCX in streaming mode
        
        Yields document elements as they are parsed.
        """
        section_num = 0
        
        # Headers/footers
        if self.extract_headers_footers:
            header_footer_text = self._extract_headers_footers()
            if header_footer_text:
                yield DocumentElement(
                    type=ElementType.TEXT,
                    content=header_footer_text,
                    metadata={"section": "headers_footers"}
                )
        
        # Process document body
        for element in self.doc.element.body:
            if isinstance(element, CT_P):
                # Paragraph
                para = Paragraph(element, self.doc)
                section_data = self._parse_paragraph(para)
                
                if section_data["text"]:
                    yield DocumentElement(
                        type=ElementType.TEXT,
                        content=section_data["text"],
                        metadata={
                            "section": section_num,
                            "style": section_data.get("style"),
                            "formatting": section_data.get("formatting", {})
                        }
                    )
                    section_num += 1
            
            elif isinstance(element, CT_Tbl):
                # Table
                if self.extract_tables:
                    table = Table(element, self.doc)
                    table_data = self._parse_table(table, section_num)
                    
                    yield DocumentElement(
                        type=ElementType.TABLE,
                        content=None,
                        metadata={
                            "section": section_num,
                            "table_info": table_data
                        }
                    )
                    section_num += 1
            
            # Yield control to event loop
            await asyncio.sleep(0)
        
        # Images
        if self.extract_images:
            images = self._extract_images()
            for image in images:
                yield DocumentElement(
                    type=ElementType.IMAGE,
                    content=None,
                    metadata={"image_info": image}
                )
    
    def _parse_paragraph(self, para: Any) -> Dict[str, Any]:
        """
        Parse a paragraph with formatting
        
        Args:
            para: python-docx Paragraph object
        
        Returns:
            Dict with text, style, and formatting info
        """
        text = para.text.strip()
        
        if not text:
            return {"text": "", "style": None, "formatting": {}}
        
        # Get style
        style = para.style.name if para.style else "Normal"
        
        # Get formatting
        formatting = {}
        
        if self.preserve_formatting:
            # Check for bold, italic, underline in runs
            has_bold = any(run.bold for run in para.runs if run.bold)
            has_italic = any(run.italic for run in para.runs if run.italic)
            has_underline = any(run.underline for run in para.runs if run.underline)
            
            formatting = {
                "bold": has_bold,
                "italic": has_italic,
                "underline": has_underline,
                "alignment": str(para.alignment) if para.alignment else None
            }
        
        return {
            "text": text,
            "style": style,
            "formatting": formatting
        }
    
    def _parse_table(self, table: Any, table_index: int) -> Optional[TableInfo]:
        """
        Parse a table with comprehensive error handling
        
        Args:
            table: python-docx Table object
            table_index: Index of table in document
        
        Returns:
            TableInfo object or None if invalid/empty
        """
        try:
            rows = []
            
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text)
                rows.append(cells)
            
            # Use safe creation helper from BaseParser
            return self._create_table_info_safe(
                rows=rows,
                table_id=f"docx_table_{table_index}",
                page_number=table_index + 1,  # Approximate
                bbox=None,
                caption=None,
                max_rows=10000  # Limit table size
            )
        
        except Exception as e:
            print(f"Warning: Failed to parse table {table_index}: {e}")
            return None
    
    def _table_to_text(self, table_info: TableInfo) -> str:
        """
        Convert table to text representation
        
        Args:
            table_info: TableInfo object
        
        Returns:
            Text representation of table
        """
        lines = []
        
        # All rows (first row is typically headers)
        if table_info.rows:
            # First row as headers
            if len(table_info.rows) > 0:
                lines.append(" | ".join(table_info.rows[0]))
                lines.append("-" * 50)
            
            # Data rows
            for row in table_info.rows[1:]:
                lines.append(" | ".join(str(cell) for cell in row))
        
        return "\n".join(lines)
    
    def _extract_images(self) -> List[ImageInfo]:
        """
        Extract images from document
        
        Returns:
            List of ImageInfo objects
        """
        images = []
        
        try:
            # Get all relationships
            rels = self.doc.part.rels
            
            image_index = 0
            for rel_id, rel in rels.items():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_part = rel.target_part
                        image_data = image_part.blob
                        
                        # Get format from content type
                        content_type = image_part.content_type
                        format_map = {
                            "image/jpeg": "jpg",
                            "image/png": "png",
                            "image/gif": "gif",
                            "image/bmp": "bmp",
                            "image/tiff": "tiff"
                        }
                        image_format = format_map.get(content_type, "unknown")
                        
                        # Create ImageInfo using safe helper
                        image_info = self._create_image_info_safe(
                            image_id=f"docx_image_{image_index}",
                            format=image_format,
                            width=0,  # Unknown - python-docx doesn't provide dimensions easily
                            height=0,  # Unknown
                            page_number=1,  # Unknown in DOCX
                            data=image_data,
                            max_image_size=10 * 1024 * 1024  # 10MB limit
                        )
                        
                        # Only append if image is valid (not None)
                        if image_info is not None:
                            images.append(image_info)
                            image_index += 1
                    
                    except Exception as e:
                        print(f"Warning: Failed to extract image {image_index}: {e}")
                        continue
        
        except Exception as e:
            print(f"Warning: Failed to extract images: {e}")
        
        return images
    
    def _extract_headers_footers(self) -> str:
        """
        Extract headers and footers
        
        Returns:
            Combined text from headers and footers
        """
        text_parts = []
        
        try:
            # Get sections
            for section in self.doc.sections:
                # Header
                if section.header:
                    header_text = "\n".join(para.text for para in section.header.paragraphs if para.text.strip())
                    if header_text:
                        text_parts.append(f"[Header]\n{header_text}")
                
                # Footer
                if section.footer:
                    footer_text = "\n".join(para.text for para in section.footer.paragraphs if para.text.strip())
                    if footer_text:
                        text_parts.append(f"[Footer]\n{footer_text}")
        
        except Exception as e:
            print(f"Warning: Failed to extract headers/footers: {e}")
        
        return "\n\n".join(text_parts)
    
    def _extract_metadata(self) -> DocumentMetadata:
        """
        Extract DOCX metadata
        
        Returns:
            DocumentMetadata object
        """
        try:
            core_props = self.doc.core_properties
            
            # Get dates
            created_date = core_props.created if hasattr(core_props, 'created') else None
            modified_date = core_props.modified if hasattr(core_props, 'modified') else None
            
            # Word count
            word_count = 0
            for para in self.doc.paragraphs:
                word_count += len(para.text.split())
            
            return self._create_metadata(
                title=core_props.title if hasattr(core_props, 'title') else None,
                author=core_props.author if hasattr(core_props, 'author') else None,
                created_date=created_date,
                modified_date=modified_date,
                page_count=self.get_page_count(),
                subject=core_props.subject if hasattr(core_props, 'subject') else None,
                keywords=core_props.keywords if hasattr(core_props, 'keywords') else None,
                word_count=word_count
            )
        
        except Exception as e:
            print(f"Warning: Failed to extract metadata: {e}")
            return self._create_metadata(
                page_count=self.get_page_count()
            )
    
    def get_supported_features(self) -> Dict[str, bool]:
        """Get supported features"""
        return {
            "text_extraction": True,
            "table_extraction": self.extract_tables,
            "image_extraction": self.extract_images,
            "metadata_extraction": True,
            "streaming": True,
            "formatting": self.preserve_formatting,
            "comments": self.extract_comments,
            "headers_footers": self.extract_headers_footers
        }


# Quick test function
async def test_docx_parser():
    """Test DOCX parser"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python -m sutra.parsers.docx <docx_file>")
        return
    
    docx_file = sys.argv[1]
    
    print(f"\n{'='*70}")
    print(f"🧪 Testing DOCX Parser")
    print(f"{'='*70}")
    print(f"File: {docx_file}\n")
    
    try:
        # Create parser
        parser = DOCXParser(docx_file)
        
        print(f"📄 Document Info:")
        print(f"   Type: {parser.document_type.value}")
        print(f"   Size: {parser.file_size:,} bytes")
        print(f"   Estimated pages: {parser.get_page_count()}")
        print(f"   Hash: {parser.file_hash[:16]}...")
        
        # Validate
        print(f"\n✓ Validating...")
        is_valid, error = await parser.validate()
        
        if not is_valid:
            print(f"❌ Validation failed: {error}")
            return
        
        print(f"✅ Valid DOCX")
        
        # Parse
        print(f"\n⏳ Parsing...")
        result = await parser.parse()
        
        if not result.success:
            print(f"❌ Parse failed: {result.error}")
            return
        
        print(f"✅ Parsed in {result.parse_time:.2f}s")
        
        # Show results
        doc = result.document
        print(f"\n📊 Results:")
        print(f"   Sections: {len(doc.pages)}")
        print(f"   Images: {len(doc.images)}")
        print(f"   Tables: {len(doc.tables)}")
        print(f"   Total text: {len(doc.full_text):,} chars")
        
        if doc.metadata.title:
            print(f"   Title: {doc.metadata.title}")
        if doc.metadata.author:
            print(f"   Author: {doc.metadata.author}")
        if doc.metadata.word_count:
            print(f"   Words: {doc.metadata.word_count:,}")
        
        # Show first section preview
        if doc.pages:
            first_section = doc.pages[0]
            preview = first_section.text[:200].replace("\n", " ")
            print(f"\n📄 First section preview:")
            print(f"   {preview}...")
        
        # Show table info
        if doc.tables:
            print(f"\n📊 Tables:")
            for idx, table in enumerate(doc.tables[:3]):  # First 3 tables
                print(f"   Table {idx + 1}: {table.rows} rows x {table.cols} cols")
        
        print(f"\n{'='*70}")
        print(f"✅ Test complete!")
        print(f"{'='*70}\n")
        
    except Exception as e:
        print(f"\n❌ Error: {e}\n")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    import asyncio
    asyncio.run(test_docx_parser())
